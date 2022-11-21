pip install streamlit
import streamlit as st
import pandas as pd
import numpy as np
#pip install pyresparser
from pyresparser import ResumeParser
import os
import spacy
#nlp = spacy.load('en_core_web_sm')
import  docx2txt
import sklearn
from sklearn.feature_extraction.text import CountVectorizer
import pickle
import base64,random
from io import StringIO
import matplotlib.pyplot as plt
import seaborn as sns
#pip install plotly
import plotly
import plotly.express as px
#import New
from plotly.subplots import make_subplots
pip install python-docx
import docx

paths = r'C:\C:\Users\91896\Desktop'
#dir_list = os.listdir(r'C:\Users\Dell\Data Science\Project - 2\Peoplesoft resumes')

st.set_page_config(
   page_title="Smart Resume Classifier",
   page_icon='Downloads/doc.png',
)

images_source="Downloads/clr2.jpg"
def set_bg_hack_url():
    st.markdown(
        f"""
         <style>
         .stApp {{
             background: url(images_source);
             background-size: cover
         }}
         </style>
         """,
        unsafe_allow_html=True
    )
set_bg_hack_url()
@st.cache(allow_output_mutation=True)
def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()
def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    page_bg_img = '''
    <style>
    body {
    background-image: url("data:image/png;base64,%s");
    background-size: cover;
    }
    </style>
    ''' % bin_str
    st.markdown(page_bg_img, unsafe_allow_html=True)
    return
set_png_as_page_bg(images_source)

st.sidebar.image('Downloads/SRA_Logo.ico')
st.title("Document Classification using NLP")
uploaded_files= st.sidebar.file_uploader("Upload Resume", type=None, accept_multiple_files=True ,key=None, help=None,
                         on_change=None, args=None, kwargs=None, disabled=False)

res_list=[]
file_name=[]
complete_path=[]
skills_list=[]


# for uploaded_file in uploaded_files:
#     def Read_files():
#         try:
#             doc = docx.Document(r'C:\Users\Dell\Data Science\Project - 2\Peoplesoft resumes\Anil kumar.docx')
#             data = ""
#             fulltext = []
#             for para in doc.paragraphs:
#                 fulltext.append(para.text)
#                 data = '\n'.join(fulltext)
#             st.write (data)
#
#         except IOError:
#             print("there was error while opening the file")
#             return
#
# if __name__ == "__main__":
#     Read_files()

def readingtextdoc(filename):
    doc = docx.Document(filename)
    completedText = []
    for paragraph in doc.paragraphs:
        completedText.append(paragraph.text)
    return '\n'.join(completedText)

def show_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    # pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf">'
    pdf_display = F'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

# Fetching Records.
for uploaded_file in uploaded_files:
    with st.spinner('Loading.....'):
        bytes_data = uploaded_file.read()
        #st.write("filename:", uploaded_file.name)
        filename=uploaded_file.name
        pfinal = os.path.join(paths, filename)
        #st.write(pfinal)
        #count1=count1+1
        complete_path.append(pfinal)
        txt = docx2txt.process(pfinal)
        if txt:
            data= txt.replace('\n',' ').replace('\t','   ')
        #return None
        res_list.append(data)
        file_name.append(filename)
        skills = ResumeParser(pfinal).get_extracted_data()
        skills_list.append(skills['skills'])
        # with open(pfinal, "wb") as f:
        #     f.write(pdf_file.getbuffer())
        #     show_pdf(pfinal)

if len(uploaded_files)>0:
    with st.spinner('Loading.....'):
        new_data = pd.DataFrame({"Resume":res_list})
        if len(uploaded_files) == 1:
            st.header(filename )
            st.subheader(readingtextdoc(pfinal))
            st.header('Skills Set - '+filename)
        else:
            st.header('Resume')
            st.write(new_data)
            st.header('Skills Set')
        # j=0
        # if j!=len(new):
        #     st.success(",".join(skills_setup[j]))
        #    #breakpoint(j)
        # else:
        #     pass
        # j=j+1
        #new = clean_data[clean_data['Resume'].isin(new_data['Resume'])]
        #new.reset_index(drop=True)
        # st.write(new['Skills_set'])
        #skills_setup = []
        #skills_setup.append(new['Skills_set'])
        k=0
        j=1
        for i in skills_list:
            if len(uploaded_files) != 1:
                st.subheader(str(j) +' . ' +((file_name)[k])+"")
            else:
                pass
            st.success(i)
            k=k+1
            j=j+1

#EDA
if len(uploaded_files)>0:
    new_data['Clean_Resume'] = new_data['Resume'].apply(lambda x: x.lower())
    new_data['Clean_Resume'] = new_data['Clean_Resume'].apply(New.remove_html_tags)
    new_data['Clean_Resume'] = new_data['Clean_Resume'].apply(New.remove_url)
    new_data['Clean_Resume'] = new_data['Clean_Resume'].apply(New.remove_punc)
    new_data['Clean_Resume'] = new_data['Clean_Resume'].apply(New.remove_stopwords)

if len(uploaded_files)>0:
    with st.spinner('Loading.....'):
        count_vector = pickle.load(open('vectorizer_document.pkl','rb'))
        model = pickle.load(open('model_document.pkl','rb'))
        vector_input = count_vector.transform(new_data['Clean_Resume'])
        #st.write(vector_input.shape)
        result = model.predict(vector_input)
        st.header('Job Designation based on Skills Set')
        #st.write(file_name)
        k=0
        count_peoplesoft=0
        count_SQl=0
        count_ReactJs=0
        count_Data_Analytics=0
        count_Software_Developer=0
        count_Others = 0
        category_values=[]
        for i in result:
            #for j in ((file_name)[k]):
            #    st.write(j)
            #st.write(i)
            if i == 1:
                st.info("Prefered  Job  Role  would  be  PeopleSoft  for  Resume - "+ ((file_name)[k]) +"")
                count_peoplesoft=count_peoplesoft+1
                category_value=1
            elif i == 2:
                st.info("Prefered  Job  Role  would  be  SQl Developer / Tester  for  Resume - "+ ((file_name)[k]) +"")
                count_SQl =count_SQl + 1
                category_value=2
            elif i == 3:
                st.info("Prefered  Job  Role  would  be  ReactJs Developer  for  Resume -  "+ ((file_name)[k]) +"")
                count_ReactJs = count_ReactJs + 1
                category_value=3
            elif i == 4:
                st.info("Prefered Job Role would be Data Analytics for Resume - "+ ((file_name)[k]) +"")
                count_Data_Analytics = count_Data_Analytics + 1
                category_value=4
            elif i == 5:
                st.info("Prefered Job Role would be Software Developer for Resume - "+ ((file_name)[k]) +"")
                count_Software_Developer = count_Software_Developer + 1
                category_value=5
            elif i == 6:
                st.info("No Job Role Found for Resume - "+ ((file_name)[k]) +"")
                count_Others = count_Others + 1
                category_value=6
            else:
                pass
            k=k+1;
            category_values.append(category_value)
        dict = {'Job': ['Peoplesoft', 'Tester or Sql Developer','React js developer','Data Analytics','Software developer','others'],
                'Values': [count_peoplesoft,count_SQl, count_ReactJs, count_Data_Analytics, count_Software_Developer,count_Others]}
        dict=pd.DataFrame(dict)

        if st.sidebar.button('Click here to check the Bar Chart'):
            mylabels = ['Peoplesoft', 'Tester or Sql Developer', 'React js developer', 'Data Analytics',
                        'Software developer', 'others']
            fig = px.bar(dict, x='Job', y='Values')
            fig.show()

        if st.sidebar.button('Click here to check the Pie Chart'):
            mylabels = ['Peoplesoft', 'Tester or Sql Developer', 'React js developer', 'Data Analytics',
                        'Software developer', 'others']
            #     # plt.pie(pie_data['Category_counts'].value_counts(),labels=list[mylabels])
            #     # plt.title('Distribution of Job title in Resume',fontsize=20)
            #     data = clean_data['Category'].value_counts()
            fig = px.pie(values=dict['Values'], names=dict['Job'])
            fig.show()

if len(uploaded_files) > 0:
    skills_list=[]
    for i in new_data['Resume']:
    #print(i)
        for j in new_data['Resume']:
            resume_data=[]
            resume_data.append(j)
            if (i == resume_data) and (j == resume_data):
            #st.write(resume_data)
                skills_list.append(new_data['Skills_set'])
        else:
            pass

    #value=[]
    #for index, item in enumerate(category_values):
        #st.write(index, item)
        #value.append(item)
    #pie_data = pd.DataFrame({'Category_counts': [value]})
    #st.write(pie_data)

    #pie_data = pd.DataFrame({'Category_counts': [category_values]})
    #st.write(pie_data)
    #value=[]
    #for i in category_values:
     #   value.append(i)
    #pie_data = pd.DataFrame({'Category_counts':[value]})
    #st.write(pie_data)
